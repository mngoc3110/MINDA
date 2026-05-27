import docx
import os

def analyze_docx_structure(file_path):
    print(f"\n--- Phân tích tệp: {os.path.basename(file_path)} ---")
    doc = docx.Document(file_path)
    
    # Đọc 20 đoạn văn đầu tiên để xem cấu trúc đề và câu hỏi
    print(" [Cấu trúc đầu đề & Câu hỏi đầu tiên]:")
    for i, para in enumerate(doc.paragraphs[:25]):
        text = para.text.strip()
        if text:
            print(f"L{i+1}: {text}")

    # Tìm xem có bảng đáp án ở cuối không
    print("\n [Kiểm tra cuối tệp - Đáp án]:")
    last_paragraphs = doc.paragraphs[-30:]
    for i, para in enumerate(last_paragraphs):
        text = para.text.strip()
        if 'đáp án' in text.lower() or 'bảng trả lời' in text.lower():
            print(f"Tìm thấy chỉ dấu đáp án: {text}")
    
    # Kiểm tra nếu có bảng (thường là bảng đáp án)
    if doc.tables:
        print(f"\n [Thông tin bảng]: Tìm thấy {len(doc.tables)} bảng trong tài liệu.")
        # Thử đọc bảng cuối cùng (thường là bảng đáp án)
        last_table = doc.tables[-1]
        print("Dữ liệu bảng cuối (trích đoạn):")
        for row in last_table.rows[:2]:
            print(" | ".join(cell.text.strip() for cell in row.cells))

folder = "projects/MINDA/de_tin_hoc"
files = [f for f in os.listdir(folder) if f.endswith('.docx')]

# Phân tích 2 tệp khác nhau để thấy sự khác biệt cấu trúc
for file in files[:2]:
    analyze_docx_structure(os.path.join(folder, file))
